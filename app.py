
import io
import json
import os
from collections import defaultdict

import streamlit as st
import pandas as pd
import plotly.express as px
from docx import Document
from docx.shared import Inches
from openai import OpenAI

try:
    import fitz
except ImportError:
    fitz = None

try:
    import cv2
    import numpy as np
except ImportError:
    cv2 = None
    np = None

CHAPTERS = [
    "Force",
    "Work, Energy and Power",
    "Machines",
    "Refraction of Light at Plane Surfaces",
    "Refraction through a Lens",
    "Spectrum",
    "Sound",
    "Current Electricity",
    "Household Circuits",
    "Electro-Magnetism",
    "Calorimetry",
    "Radioactivity",
]

MODEL = os.getenv("OPENAI_MODEL", "gpt-5.6-luna")

st.set_page_config(page_title="AI Question Paper Analyzer", page_icon="📘", layout="wide")

st.markdown("""
<style>
.block-container{padding-top:1.4rem;padding-bottom:3rem}
.hero{padding:1.2rem 1.3rem;border:1px solid #ddd;border-radius:16px;
background:linear-gradient(135deg,rgba(139,47,201,.08),rgba(255,255,255,.85))}
.hero h1{margin:0}.muted{color:#666}
.chapter-title{font-size:1.18rem;font-weight:750;border-bottom:2px solid #ddd;
padding-bottom:.4rem;margin-top:1rem}
.tag{display:inline-block;padding:.25rem .65rem;border-radius:999px;
background:#f3e8ff;color:#7a27b0;font-weight:650;font-size:.84rem}
</style>
""", unsafe_allow_html=True)


def client():
    key = os.getenv("OPENAI_API_KEY")
    if not key:
        try:
            key = st.secrets["OPENAI_API_KEY"]
        except Exception:
            key = None
    if not key:
        st.error("OPENAI_API_KEY is missing. Add it in Streamlit Secrets.")
        st.stop()
    return OpenAI(api_key=key)


def docx_text(data):
    doc = Document(io.BytesIO(data))
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            vals = [c.text.strip() for c in row.cells]
            if any(vals):
                parts.append(" | ".join(vals))
    return "\n".join(parts)


def upload_input(upload):
    data = upload.getvalue()
    name = upload.name.lower()
    if name.endswith(".docx"):
        return [{"type":"input_text","text":docx_text(data)}]
    if name.endswith(".txt"):
        return [{"type":"input_text","text":data.decode("utf-8",errors="ignore")}]
    f = client().files.create(file=(upload.name, data), purpose="user_data")
    return [{"type":"input_file","file_id":f.id}]


def output_schema():
    return {
        "type":"object","additionalProperties":False,
        "properties":{
            "paper_title":{"type":"string"},
            "subject":{"type":"string"},
            "printed_question_marks":{"type":"number"},
            "maximum_exam_marks":{"type":"number"},
            "section_structure":{"type":"string"},
            "questions":{
                "type":"array","items":{
                    "type":"object","additionalProperties":False,
                    "properties":{
                        "question_no":{"type":"string"},
                        "parent_question":{"type":"string"},
                        "section":{"type":"string"},
                        "question_text":{"type":"string"},
                        "chapter":{"type":"string","enum":CHAPTERS},
                        "marks":{"type":"number"},
                        "concept":{"type":"string"},
                        "question_type":{"type":"string"},
                        "difficulty":{"type":"string","enum":["Easy","Moderate","Difficult","Not clear"]},
                        "split_reason":{"type":"string"},
                        "source_page":{"type":"integer"},
                        "has_figure":{"type":"boolean"},
                        "figure_x1":{"type":"number"},
                        "figure_y1":{"type":"number"},
                        "figure_x2":{"type":"number"},
                        "figure_y2":{"type":"number"},
                    },
                    "required":[
                        "question_no","parent_question","section","question_text",
                        "chapter","marks","concept","question_type","difficulty",
                        "split_reason","source_page","has_figure",
                        "figure_x1","figure_y1","figure_x2","figure_y2"
                    ]
                }
            },
            "focus_concepts":{
                "type":"array","items":{
                    "type":"object","additionalProperties":False,
                    "properties":{
                        "concept":{"type":"string"},
                        "chapter":{"type":"string","enum":CHAPTERS},
                        "reason":{"type":"string"}
                    },
                    "required":["concept","chapter","reason"]
                }
            }
        },
        "required":[
            "paper_title","subject","printed_question_marks","maximum_exam_marks",
            "section_structure","questions","focus_concepts"
        ]
    }


def analyse(upload):
    chapter_list = "\n".join(f"{i+1}. {x}" for i,x in enumerate(CHAPTERS))
    prompt = f"""
You are an expert school Physics examination-paper analyst.

Allowed chapters ONLY:
{chapter_list}

Analyse the complete uploaded question paper.

EXTRACTION:
- Extract every assessable question and preserve its original numbering.
- Split independently assessable (a), (b), (c) parts.
- If a sub-question contains distinct assessable parts from different chapters,
  split those further.
- Keep each extracted question's text limited to THAT question only. Do not
  accidentally append text from the preceding or following question.

MARKS:
- printed_question_marks = total marks printed across the complete paper,
  including all optional questions.
- maximum_exam_marks = maximum score a student can obtain after applying all
  attempt/choice rules.
- Do not confuse printed total with exam maximum.
- For the benchmark paper, Section A is 40 compulsory, Section B prints six
  10-mark questions, any four are attempted, so printed total is 100 and
  maximum exam score is 80.

FIGURES:
- Identify source_page for each question.
- If a question depends on a diagram, graph, circuit, ray diagram, apparatus,
  table, or other visual, set has_figure=true.
- Give a TIGHT normalized bounding box (0..1) around ONLY the visual needed
  for that question.
- Include the visual's own labels, arrows, axes and legend.
- EXCLUDE question prose, answer choices, marks, headings and adjacent-question
  text. Do not use a large box merely because it is convenient.
- If no visual, has_figure=false and coordinates all 0.

CLASSIFICATION:
- Assign exactly one allowed chapter to every extracted item.
- Identify concept, type and difficulty.
- Give concise reasons for any split.
- Identify important focus concepts.

Return structured JSON only.
"""
    r = client().responses.create(
        model=MODEL,
        input=[{"role":"user","content":upload_input(upload)+[{"type":"input_text","text":prompt}]}],
        text={"format":{"type":"json_schema","name":"physics_qp_v7","strict":True,"schema":output_schema()}}
    )
    return json.loads(r.output_text)


def answer(q):
    prompt = f"""
Answer this school Physics question as a teacher.
Chapter: {q['chapter']}
Concept: {q['concept']}
Marks: {q['marks']}
Question: {q['question_text']}

Match the depth to the marks. For numericals use Given, Formula,
Substitution, Calculation and Final Answer. Keep units. For theory give
clear exam-ready points. If a diagram is essential, describe it.
"""
    return client().responses.create(model=MODEL,input=prompt).output_text.strip()



def section_mark_summary(data):
    """Summarize printed marks by section using extracted question numbers."""
    sections = defaultdict(float)
    for q in data["questions"]:
        sections[q.get("section", "Other")] += float(q.get("marks", 0))
    return dict(sections)


def benchmark_mark_validation(data):
    """
    The benchmark paper has a known structure:
    Section A = 40 compulsory
    Section B = six printed 10-mark questions = 60
    Printed total = 100
    Actual maximum = 80

    Return a diagnostic instead of silently changing the AI's extracted marks.
    """
    sections = section_mark_summary(data)
    # Section labels vary by OCR/AI, so also use the question numbering when needed.
    a = sum(float(q["marks"]) for q in data["questions"]
            if str(q.get("question_no","")).split("(")[0].strip().upper() in
            {"1","2","3"})
    b = sum(float(q["marks"]) for q in data["questions"]
            if str(q.get("question_no","")).split("(")[0].strip().upper() in
            {"4","5","6","7","8","9"})
    return {
        "section_a_extracted": a,
        "section_b_extracted": b,
        "expected_printed": 100.0,
        "expected_exam_max": 80.0,
        "extracted_total": a + b,
        "difference": round((a + b) - 100.0, 2),
    }



def chapter_totals(data):
    """Return chapter-wise marks and question counts for the analysed paper."""
    totals={c:{"marks":0.0,"questions":0} for c in CHAPTERS}
    for q in data.get("questions",[]):
        c=q.get("chapter")
        if c not in totals:
            continue
        totals[c]["marks"] += float(q.get("marks",0) or 0)
        totals[c]["questions"] += 1
    return totals


def figure_crop(pdf_bytes, page_no, coords):
    """
    Crop a question figure conservatively.

    Priority:
    1. Embedded/vector visual objects inside the AI box.
    2. A conservative AI box for scanned pages.

    We deliberately avoid aggressive contour grouping because it can merge
    nearby question text into the diagram. If no reliable visual object exists,
    returning the AI crop is safer than inventing a crop.
    """
    if fitz is None:
        return None

    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        if not (1 <= int(page_no) <= len(doc)):
            doc.close()
            return None

        page = doc[int(page_no) - 1]
        p = page.rect
        vals = [max(0.0, min(1.0, float(v))) for v in coords]
        candidate = fitz.Rect(
            p.x0 + vals[0] * p.width,
            p.y0 + vals[1] * p.height,
            p.x0 + vals[2] * p.width,
            p.y0 + vals[3] * p.height,
        )

        # Guard against malformed AI coordinates.
        if candidate.width <= 3 or candidate.height <= 3:
            doc.close()
            return None

        visual_boxes = []

        # Vector drawings.
        try:
            for d in page.get_drawings():
                r = d.get("rect")
                if not r or r.width <= 1 or r.height <= 1:
                    continue
                if not r.intersects(candidate):
                    continue
                inter = r & candidate
                # Ignore tiny marks.
                if inter.width * inter.height >= 12:
                    visual_boxes.append(r)
        except Exception:
            pass

        # Embedded raster images.
        try:
            for item in page.get_image_info(xrefs=True):
                r = item.get("bbox")
                if not r:
                    continue
                r = fitz.Rect(r)
                if r.intersects(candidate) and r.width * r.height >= 150:
                    visual_boxes.append(r)
        except Exception:
            pass

        crop = candidate

        if visual_boxes:
            # Use objects whose centres are inside the AI candidate.
            inside = []
            for r in visual_boxes:
                cx = (r.x0 + r.x1) / 2
                cy = (r.y0 + r.y1) / 2
                if candidate.contains(fitz.Point(cx, cy)):
                    inside.append(r)

            if inside:
                # Keep only a compact cluster around the candidate centre.
                ccx = (candidate.x0 + candidate.x1) / 2
                ccy = (candidate.y0 + candidate.y1) / 2
                inside.sort(
                    key=lambda r:
                    ((r.x0+r.x1)/2-ccx)**2 + ((r.y0+r.y1)/2-ccy)**2
                )

                # Start with nearest object, then include nearby objects that
                # are plausibly part of the same diagram.
                selected = [inside[0]]
                for r in inside[1:]:
                    current = selected[-1]
                    gap_x = max(0, max(current.x0, r.x0) - min(current.x1, r.x1))
                    gap_y = max(0, max(current.y0, r.y0) - min(current.y1, r.y1))
                    if gap_x < candidate.width * 0.12 and gap_y < candidate.height * 0.18:
                        selected.append(r)

                crop = selected[0]
                for r in selected[1:]:
                    crop |= r

                # Small margin to retain labels/arrows.
                mx = min(8, p.width * 0.008)
                my = min(8, p.height * 0.008)
                crop = fitz.Rect(
                    max(p.x0, crop.x0-mx),
                    max(p.y0, crop.y0-my),
                    min(p.x1, crop.x1+mx),
                    min(p.y1, crop.y1+my),
                )

        pix = page.get_pixmap(matrix=fitz.Matrix(2.5, 2.5), clip=crop, alpha=False)
        png = pix.tobytes("png")
        doc.close()
        return png

    except Exception:
        return None


def make_word(data,pdf_bytes=None,with_answers=False):
    doc=Document()
    doc.add_heading(data.get("paper_title") or "AI Sorted Question Paper",0)
    doc.add_paragraph(f"Subject: {data.get('subject','')}")
    doc.add_paragraph(
        f"Printed paper: {data.get('printed_question_marks',100):g} marks | "
        f"Actual exam maximum: {data.get('maximum_exam_marks',80):g} marks"
    )
    t=chapter_totals(data); total=sum(v["marks"] for v in t.values())
    doc.add_heading("Chapter-wise Weightage",1)
    doc.add_paragraph(f"Weightage is calculated from the complete printed paper ({total:g} marks).")
    table=doc.add_table(rows=1,cols=4)
    for cell,label in zip(table.rows[0].cells,["Chapter","Questions","Marks","Weightage"]):
        cell.text=label
    for c in CHAPTERS:
        if t[c]["marks"]:
            cells=table.add_row().cells; m=t[c]["marks"]
            cells[0].text=c; cells[1].text=str(t[c]["questions"])
            cells[2].text=str(m); cells[3].text=f"{m/total*100:.1f}%"

    groups=defaultdict(list)
    for q in data["questions"]: groups[q["chapter"]].append(q)
    doc.add_page_break(); doc.add_heading("Questions Sorted Chapter-wise",1)
    for c in CHAPTERS:
        if not groups[c]: continue
        doc.add_heading(c,2)
        for q in groups[c]:
            p=doc.add_paragraph(); p.add_run(f"Q{q['question_no']}. ").bold=True
            p.add_run(q["question_text"])
            meta=doc.add_paragraph()
            meta.add_run("Marks: ").bold=True; meta.add_run(str(q["marks"]))
            meta.add_run(" | Concept: ").bold=True; meta.add_run(q["concept"])
            if pdf_bytes and q.get("has_figure"):
                img=figure_crop(pdf_bytes,q.get("source_page",0),
                                (q.get("figure_x1",0),q.get("figure_y1",0),
                                 q.get("figure_x2",0),q.get("figure_y2",0)))
                if img:
                    doc.add_paragraph("Figure:")
                    doc.add_picture(io.BytesIO(img),width=Inches(5.8))
            if with_answers:
                h=doc.add_paragraph(); h.add_run("AI Answer").bold=True
                doc.add_paragraph(answer(q))
    out=io.BytesIO(); doc.save(out); return out.getvalue()


def reset():
    for k in ("data","answers","source_bytes","source_name"):
        st.session_state[k]=None if k=="data" or k=="source_bytes" or k=="source_name" else {}


if "data" not in st.session_state: st.session_state.data=None
if "answers" not in st.session_state: st.session_state.answers={}
if "source_bytes" not in st.session_state: st.session_state.source_bytes=None
if "source_name" not in st.session_state: st.session_state.source_name=None

st.markdown("""
<div class="hero">
<h1>📘 AI Question Paper Analyzer</h1>
<div class="muted">Upload → classify → analyse → focus → answer → download.</div>
</div>
""",unsafe_allow_html=True)

with st.sidebar:
    st.header("📚 12-Chapter Syllabus")
    for i,c in enumerate(CHAPTERS,1): st.write(f"**{i}.** {c}")
    st.divider(); st.caption(f"AI model: {MODEL}")
    if st.session_state.data and st.button("🆕 Start New Paper",use_container_width=True):
        reset(); st.rerun()

upload=st.file_uploader("Upload Question Paper",
                        type=["pdf","docx","txt","png","jpg","jpeg","webp"])
if st.button("🔍 Analyse Paper",type="primary",disabled=upload is None,use_container_width=True):
    st.session_state.data=None; st.session_state.answers={}
    st.session_state.source_bytes=upload.getvalue(); st.session_state.source_name=upload.name
    with st.spinner("Reading questions, locating figures and classifying chapters..."):
        try: st.session_state.data=analyse(upload); st.success("Analysis complete.")
        except Exception as e: st.error(f"Analysis failed: {e}")

data=st.session_state.data
if not data:
    st.info("Upload a paper to begin."); st.stop()

st.info(data["section_structure"])
t=chapter_totals(data)
printed=sum(v["marks"] for v in t.values())
exam=float(data.get("maximum_exam_marks") or 80)

a,b,c,d=st.columns(4)
a.metric("Assessable items",len(data["questions"]))
b.metric("Printed paper",f"{printed:g}")
c.metric("Actual exam maximum",f"{exam:g}")
d.metric("Weightage total","100%")

validation = benchmark_mark_validation(data)
if abs(validation["difference"]) > 0.01:
    st.warning(
        f"⚠️ Mark reconciliation: the extracted questions total "
        f"{validation['extracted_total']:g} marks, which differs from the "
        f"expected 100-mark printed paper by {validation['difference']:+g}. "
        "Review the question marks below before relying on the weightage chart."
    )

st.header("📊 Chapter-wise Weightage")
st.caption(
    f"Chapter weightage uses the complete printed paper ({printed:g} marks). "
    f"The actual student maximum is {exam:g} marks because of the Section B choice."
)
rows=[{"Chapter":c,"Questions":t[c]["questions"],"Marks":t[c]["marks"],
       "Weightage %":round(t[c]["marks"]/printed*100,1)}
      for c in CHAPTERS if t[c]["marks"]]
if rows:
    df=pd.DataFrame(rows)
    fig=px.pie(df,names="Chapter",values="Marks",hole=.4,
               title="Chapter Weightage — Complete Printed Paper")
    fig.update_traces(textposition="inside",textinfo="percent")
    st.plotly_chart(fig,use_container_width=True)
    view=df.copy(); view["Weightage %"]=view["Weightage %"].astype(str)+"%"
    st.dataframe(view,use_container_width=True,hide_index=True)

st.header("🎯 Main Concepts to Focus On")
for x in data["focus_concepts"]:
    st.markdown(f"**{x['concept']}** — `{x['chapter']}`  \n{x['reason']}")

st.header("✏️ Teacher Review")
st.caption("Change any chapter assignment. The sorted sections and weightage update.")
for i,q in enumerate(data["questions"]):
    with st.container(border=True):
        l,r=st.columns([5,2])
        with l:
            st.markdown(f"**Q{q['question_no']}.** {q['question_text']}")
            st.markdown(f'<span class="tag">Chapter: {q["chapter"]}</span>',unsafe_allow_html=True)
            st.caption(f"Marks: {q['marks']} • {q['concept']} • {q['question_type']} • {q['difficulty']}")
            if q.get("has_figure"): st.caption(f"📐 Figure detected — page {q.get('source_page','?')}")
        with r:
            choice=st.selectbox("Chapter",CHAPTERS,index=CHAPTERS.index(q["chapter"]),key=f"ch_{i}")
            if choice!=q["chapter"]:
                data["questions"][i]["chapter"]=choice; st.rerun()

st.header("📝 Questions Sorted Chapter-wise")
groups=defaultdict(list)
for i,q in enumerate(data["questions"]): groups[q["chapter"]].append((i,q))

for c in CHAPTERS:
    if not groups[c]: continue
    st.markdown(f'<div class="chapter-title">{c}</div>',unsafe_allow_html=True)
    for i,q in groups[c]:
        with st.container(border=True):
            st.markdown(f"**Q{q['question_no']}.** {q['question_text']}")
            st.markdown(f'<span class="tag">Chapter: {q["chapter"]}</span>',unsafe_allow_html=True)
            x,y,z=st.columns(3)
            x.caption(f"Marks: {q['marks']} • {q['question_type']}")
            y.caption(f"Concept: {q['concept']}")
            z.caption(f"Difficulty: {q['difficulty']}")
            if q.get("has_figure"):
                img=None
                if (st.session_state.source_bytes and
                    st.session_state.source_name.lower().endswith(".pdf")):
                    img=figure_crop(st.session_state.source_bytes,q.get("source_page",0),
                                    (q.get("figure_x1",0),q.get("figure_y1",0),
                                     q.get("figure_x2",0),q.get("figure_y2",0)))
                if img: st.image(img,caption="Original figure (cropped)")
                else: st.caption("📐 Original figure is referenced in the Word export if available.")
            key=f"{q['question_no']}::{q['question_text']}"
            if key in st.session_state.answers:
                with st.expander("💡 AI Answer",expanded=True): st.markdown(st.session_state.answers[key])
            elif st.button("🔗 View AI Answer",key=f"ans_{i}"):
                with st.spinner("Generating answer..."):
                    try:
                        st.session_state.answers[key]=answer(q); st.rerun()
                    except Exception as e: st.error(f"Answer generation failed: {e}")

st.divider(); st.header("⬇️ Download")
x,y=st.columns(2)
with x:
    st.download_button("📄 Sorted Question Paper (Word)",
        data=make_word(data,st.session_state.source_bytes,False),
        file_name="sorted_question_paper.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        type="primary",use_container_width=True)
with y:
    st.download_button("📄 Sorted Paper + AI Answers",
        data=make_word(data,st.session_state.source_bytes,True),
        file_name="sorted_question_paper_with_answers.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True)
